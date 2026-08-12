from io import BytesIO

from docx import Document

from app.domain.audience_persona import VerifiedAudiencePersona
from app.domain.claim import VerifiedClaim
from app.domain.market_research_document import MarketResearchDocument
from app.domain.section_result import SectionResult
from app.export.docx_builder import build_docx


def _claim(text: str, section: str, group_key: str | None = None, field_key: str | None = None) -> VerifiedClaim:
    return VerifiedClaim(
        claim_id=f"c-{text}", section=section, text=text, chunk_id="c1", block_span=(0, 0),
        verified=True, group_key=group_key, field_key=field_key,
    )


def _all_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _minimal_document(sections: dict[str, SectionResult]) -> MarketResearchDocument:
    return MarketResearchDocument(
        id="doc-1", brand_id="acme-roasters", status="pending_approval", sections=sections, call_site_trace={}
    )


def test_prose_section_renders_as_paragraphs():
    doc = _minimal_document(
        {"brand_overview": SectionResult(section="brand_overview", brand_id="acme-roasters", status="verified", claims=[_claim("Acme sells specialty coffee.", "brand_overview")])}
    )
    document = Document(build_docx(doc))
    assert "Acme sells specialty coffee." in _all_text(document)


def test_competitor_table_groups_claims_into_rows_and_columns():
    claims = [
        _claim("Fast shipping", "competitor_analysis", group_key="Acme Corp", field_key="strengths"),
        _claim("Slow support", "competitor_analysis", group_key="Acme Corp", field_key="weaknesses"),
        _claim("Cheap pricing", "competitor_analysis", group_key="Beta Inc", field_key="strengths"),
    ]
    doc = _minimal_document(
        {"competitor_analysis": SectionResult(section="competitor_analysis", brand_id="acme-roasters", status="verified", claims=claims)}
    )
    document = Document(build_docx(doc))
    text = _all_text(document)
    assert "Acme Corp" in text
    assert "Beta Inc" in text
    assert "Fast shipping" in text
    assert "Offer & positioning" in text  # header row present even though unused here


def test_platform_table_renders_all_seven_fixed_rows_even_with_no_data():
    doc = _minimal_document(
        {"platform_analysis": SectionResult(section="platform_analysis", brand_id="acme-roasters", status="insufficient_evidence", claims=[])}
    )
    document = Document(build_docx(doc))
    text = _all_text(document)
    for platform in ["Instagram", "Facebook", "LinkedIn", "YouTube", "X (Twitter)", "Threads", "Pinterest"]:
        assert platform in text


def test_swot_grid_groups_claims_into_four_quadrants():
    claims = [
        _claim("DISCOM track record", "swot", field_key="strength"),
        _claim("Dated visual identity", "swot", field_key="weakness"),
        _claim("Smart-grid market growth", "swot", field_key="opportunity"),
        _claim("New entrants undercutting price", "swot", field_key="threat"),
    ]
    doc = _minimal_document({"swot": SectionResult(section="swot", brand_id="acme-roasters", status="verified", claims=claims)})
    document = Document(build_docx(doc))
    # platform_analysis's fixed-row table always renders too (even with no
    # data for this fixture doc) -- find the SWOT table specifically rather
    # than assuming table order.
    swot_table = next(t for t in document.tables if "Strengths" in t.rows[0].cells[0].text)
    assert len(swot_table.rows) == 2
    assert len(swot_table.columns) == 2
    text = _all_text(document)
    for label in ["Strengths", "Weaknesses", "Opportunities", "Threats"]:
        assert label in text
    assert "DISCOM track record" in text


def test_persona_table_transposes_attributes_as_rows_personas_as_columns():
    personas = [
        VerifiedAudiencePersona(
            persona_id="p1", section="target_audience", name="Weekend warrior",
            pain_points=["Limited time"], interests=["Quality gear"], chunk_ids=["c1"],
            age_range="25-34", location="Urban US", occupation_income="Mid-career",
            preferred_platforms=["Instagram"], verified=True,
        )
    ]
    doc = _minimal_document(
        {"target_audience": SectionResult(section="target_audience", brand_id="acme-roasters", status="verified", personas=personas)}
    )
    document = Document(build_docx(doc))
    text = _all_text(document)
    assert "Persona 1" in text
    assert "Name / label" in text
    assert "Weekend warrior" in text
    assert "25-34" in text


def test_missing_section_degrades_to_honest_placeholder_not_a_crash():
    doc = _minimal_document({})  # no sections at all
    buffer: BytesIO = build_docx(doc)
    document = Document(buffer)
    assert "No verified findings yet." in _all_text(document)
