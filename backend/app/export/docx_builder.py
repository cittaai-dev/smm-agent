"""Renders a MarketResearchDocument into a .docx matching
TEMPLATE_1_Market_Research.docx's section-by-section structure -- a pure,
computed projection over already-verified data (same "never a second write
path" discipline as domain/client_view.py), grouping claims on
claim.group_key/field_key exactly like the frontend's table/grid components
do. No LLM involved -- this is presentation, not generation."""

from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

from app.domain.audience_persona import VerifiedAudiencePersona
from app.domain.claim import VerifiedClaim
from app.domain.market_research_document import MarketResearchDocument
from app.domain.sop1 import SOP1_SECTIONS, SectionSpec

_HEADER_FILL = "3EBD93"  # teal, matches TEMPLATE_1_Market_Research.docx's table headers

_FIELD_LABELS = {
    "offer_positioning": "Offer & positioning",
    "strengths": "Strengths",
    "weaknesses": "Weaknesses",
    "content_frequency": "Content & frequency",
    "gaps_to_use": "Gaps to use",
    "audience_here": "Audience here? (Yes/No)",
    "priority": "Priority (High/Med/Low)",
    "notes": "Notes",
}

_PERSONA_ROWS = [
    ("name", "Name / label"),
    ("age_range", "Age range"),
    ("location", "Location"),
    ("occupation_income", "Occupation / income"),
    ("interests", "Interests"),
    ("pain_points", "Pain points"),
    ("preferred_platforms", "Preferred platforms"),
]

_SWOT_QUADRANTS = [
    ("strength", "Strengths", "DCFCE7"),
    ("weakness", "Weaknesses", "FEE2E2"),
    ("opportunity", "Opportunities", "EDE9FE"),
    ("threat", "Threats", "FFEDD5"),
]


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell(cell: _Cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.text = text
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].bold = bold
    if fill:
        _shade_cell(cell, fill)


def _group_claims(claims: list[VerifiedClaim]) -> dict[str, dict[str, str]]:
    """group_key -> field_key -> text, verified claims only -- an unverified
    or untagged claim simply isn't part of the table (P4 applies to the
    export the same as the live document view)."""
    groups: dict[str, dict[str, str]] = {}
    for c in claims:
        if c.verified and c.group_key and c.field_key:
            groups.setdefault(c.group_key, {})[c.field_key] = c.text
    return groups


def _add_prose(document: Document, claims: list[VerifiedClaim]) -> None:
    verified = [c for c in claims if c.verified]
    if not verified:
        document.add_paragraph("No verified findings yet.")
        return
    for claim in verified:
        document.add_paragraph(claim.text)


def _add_persona_table(document: Document, personas: list[VerifiedAudiencePersona]) -> None:
    verified = [p for p in personas if p.verified]
    if not verified:
        document.add_paragraph("No verified personas yet.")
        return
    table = document.add_table(rows=len(_PERSONA_ROWS) + 1, cols=len(verified) + 1)
    table.style = "Table Grid"
    _set_cell(table.rows[0].cells[0], "Persona detail", bold=True, fill=_HEADER_FILL)
    for col, persona in enumerate(verified, start=1):
        _set_cell(table.rows[0].cells[col], f"Persona {col}", bold=True, fill=_HEADER_FILL)
    for row, (field, label) in enumerate(_PERSONA_ROWS, start=1):
        _set_cell(table.rows[row].cells[0], label, bold=True)
        for col, persona in enumerate(verified, start=1):
            value = getattr(persona, field)
            if isinstance(value, list):
                value = ", ".join(value)
            table.rows[row].cells[col].text = value or ""


def _add_grouped_table(
    document: Document, claims: list[VerifiedClaim], spec: SectionSpec, row_header: str
) -> None:
    groups = _group_claims(claims)
    fields = spec.structured_fields
    row_keys = spec.structured_row_values or list(groups.keys())
    if not any(row_keys):
        document.add_paragraph("No verified findings yet.")
        return
    table = document.add_table(rows=1, cols=len(fields) + 1)
    table.style = "Table Grid"
    _set_cell(table.rows[0].cells[0], row_header, bold=True, fill=_HEADER_FILL)
    for col, field in enumerate(fields, start=1):
        _set_cell(table.rows[0].cells[col], _FIELD_LABELS.get(field, field), bold=True, fill=_HEADER_FILL)
    for row_key in row_keys:
        row = table.add_row().cells
        row[0].text = row_key
        field_map = groups.get(row_key, {})
        for col, field in enumerate(fields, start=1):
            row[col].text = field_map.get(field, "")


def _add_swot_grid(document: Document, claims: list[VerifiedClaim]) -> None:
    buckets: dict[str, list[str]] = {key: [] for key, _, _ in _SWOT_QUADRANTS}
    for claim in claims:
        if claim.verified and claim.field_key in buckets:
            buckets[claim.field_key].append(claim.text)
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for i, (key, label, color) in enumerate(_SWOT_QUADRANTS):
        cell = table.rows[i // 2].cells[i % 2]
        _set_cell(cell, label, bold=True, fill=color)
        items = buckets[key]
        if items:
            for item in items:
                cell.add_paragraph(f"• {item}")
        else:
            cell.add_paragraph("No verified findings yet.")


def build_docx(doc: MarketResearchDocument) -> BytesIO:
    document = Document()
    document.add_heading(f"{doc.brand_id} — Market Research", level=0)

    for n, spec in enumerate(SOP1_SECTIONS, start=1):
        section = doc.sections.get(spec.id)
        document.add_heading(f"{n}. {spec.label}", level=1)

        # platform_table has a fixed row set (the 7 platforms) independent of
        # whether any claim exists yet -- it must still render its full
        # skeleton, not fall into the generic "nothing here" placeholder.
        has_content = section is not None and (section.claims or section.personas)
        if not has_content and spec.structured_output != "platform_table":
            document.add_paragraph("No verified findings yet.")
            continue

        claims = section.claims if section is not None else []
        personas = section.personas if section is not None else []

        if spec.id == "target_audience":
            _add_prose(document, claims)
            _add_persona_table(document, personas)
        elif spec.structured_output == "competitor_table":
            _add_grouped_table(document, claims, spec, row_header="Competitor")
        elif spec.structured_output == "platform_table":
            _add_grouped_table(document, claims, spec, row_header="Platform")
        elif spec.structured_output == "swot_grid":
            _add_swot_grid(document, claims)
        else:
            _add_prose(document, claims)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
